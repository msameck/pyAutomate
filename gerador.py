import gspread
from oauth2client.service_account import ServiceAccountCredentials
from docx import Document
from datetime import datetime

credenciais_json = "caminho_para_suas_credenciais.json"

scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
creds = ServiceAccountCredentials.from_json_keyfile_name(credenciais_json, scope)
client = gspread.authorize(creds)

spreadsheet = client.open_by_url('https://docs.google.com/spreadsheets/d/1oqgIr-p0xo9rT2KTkmlhQDpdvXR7YG1JvK2ITYetGuI/edit?usp=sharing')
sheet = spreadsheet.sheet1

data = sheet.get_all_records()

hoje = datetime.today().strftime('%Y-%m-%d')

servico_map = {
    'ABONO DE FALTA': 'ab_falta',
    'ABONO FAMILIAR': 'ab_familiar',
    'ANOTAÇÃO EM MINHA FICHA FUNCIONAL': 'aeff',
    'DECLARAÇÃO': 'dc',
    'DIVERSOS': 'dv',
    'EXONERAÇÃO': 'ex',
    'FÊRIAS': 'frs',
    'LICENÇA MATERNIDADE': 'lm',
    'LICENÇA PATERNIDADE': 'lp',
    'OUTRO TIPO DE LICENÇA': 'ou'
}

for idx, row in enumerate(data):
    data_adicao = row['CARIMBO DE DATA/HORA'].split()[0]
    processado = row.get('Processado', 'Não')

    if data_adicao == hoje and processado != 'Sim':
        doc = Document("modelo.docx")

        for paragraph in doc.paragraphs:
            if '{CARIMBO DE DATA/HORA}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{CARIMBO DE DATA/HORA}', row['CARIMBO DE DATA/HORA'])
            if '{NOME COMPLETO}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{NOME COMPLETO}', row['NOME COMPLETO'])
            if '{ESTADO CIVIL}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{ESTADO CIVIL}', row['ESTADO CIVIL'])
            if '{ENDEREÇO}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{ENDEREÇO}', row['ENDEREÇO'])
            if '{NÚMERO DA RESIDÊNCIA}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{NÚMERO DA RESIDÊNCIA}', row['NÚMERO DA RESIDÊNCIA'])
            if '{CIDADE}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{CIDADE}', row['CIDADE'])
            if '{BAIRRO}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{BAIRRO}', row['BAIRRO'])
            if '{TELEFONE / CELULAR}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{TELEFONE / CELULAR}', row['TELEFONE / CELULAR'])
            if '{EMAIL}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{EMAIL}', row['EMAIL'])
            if '{CPF}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{CPF}', row['CPF'])
            if '{LOCAL DE TRABALHO}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{LOCAL DE TRABALHO}', row['LOCAL DE TRABALHO'])
            if '{MATRÍCULA}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{MATRÍCULA}', row['MATRÍCULA'])
            if '{CARGO}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{CARGO}', row['CARGO'])
            if '{SERVIÇO DESEJADO}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{SERVIÇO DESEJADO}', row['SERVIÇO DESEJADO'])
            if '{OUTROS - ESPECIFICAR ASSUNTO}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{OUTROS - ESPECIFICAR ASSUNTO}', row['OUTROS - ESPECIFICAR ASSUNTO'])

        servico_desejado = row['SERVIÇO DESEJADO']
        for key in servico_map:
            placeholder = servico_map[key]
            if key == servico_desejado:
                for paragraph in doc.paragraphs:
                    if f'{{{placeholder}}}' in paragraph.text:
                        paragraph.text = paragraph.text.replace(f'{{{placeholder}}}', 'X')
            else:
                for paragraph in doc.paragraphs:
                    if f'{{{placeholder}}}' in paragraph.text:
                        paragraph.text = paragraph.text.replace(f'{{{placeholder}}}', ' ')

        doc.save(f"{row['NOME COMPLETO']}.docx")

        sheet.update_cell(idx + 2, sheet.find('Processado').col, 'Sim')
